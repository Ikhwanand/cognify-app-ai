from sqlmodel import Session, select, delete
from models.document import Document, DocumentChunk
from pypdf import PdfReader
from docx import Document as DocxDocument
from typing import List, BinaryIO


class DocumentService:
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def extract_text(self, file: BinaryIO, filename: str) -> str:
        """Extract text from various file types"""
        file_ext = filename.lower().split(".")[-1]

        if file_ext == "pdf":
            return self._extract_pdf(file)
        elif file_ext == "docx":
            return self._extract_docx(file)
        elif file_ext in ["txt", "md"]:
            return file.read().decode("utf-8")
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")

    def _extract_pdf(self, file: BinaryIO) -> str:
        reader = PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text

    def _extract_docx(self, file: BinaryIO) -> str:
        doc = DocxDocument(file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text

    def chunk_text(self, text: str) -> List[str]:
        """Split text into chunks with overlap"""
        if not text:
            return []

        chunks = []
        start = 0

        while start < len(text):
            end = start + self.chunk_size
            chunk = text[start:end]

            if end < len(text):
                last_period = chunk.rfind(".")
                if last_period > self.chunk_size * 0.5:
                    chunk = chunk[: last_period + 1]
                    end = start + len(chunk)

            chunks.append(chunk.strip())
            start = end - self.chunk_overlap

        return [c for c in chunks if c]

    async def process_document(
        self, db: Session, file: BinaryIO, filename: str, file_size: int
    ) -> Document:
        """Process and store a document"""
        content = self.extract_text(file, filename)
        file_ext = filename.lower().split(".")[-1]

        # Create document
        document = Document(
            name=filename, file_type=file_ext, file_size=file_size, content=content
        )
        db.add(document)
        db.commit()
        db.refresh(document)

        # Create chunks
        chunks = self.chunk_text(content)
        for i, chunk_content in enumerate(chunks):
            chunk = DocumentChunk(
                document_id=document.id, content=chunk_content, chunk_index=i
            )
            db.add(chunk)

        document.chunk_count = len(chunks)
        db.commit()
        db.refresh(document)

        return document

    def get_all_documents(self, db: Session) -> List[Document]:
        statement = select(Document).order_by(Document.created_at.desc())
        return db.exec(statement).all()

    def delete_document(self, db: Session, document_id: str) -> bool:
        """Delete a document and all its chunks"""
        document = db.get(Document, document_id)
        if document:
            # Delete chunks first using direct DELETE statement
            db.exec(
                delete(DocumentChunk).where(DocumentChunk.document_id == document_id)
            )
            db.commit()

            # Now delete the document
            db.delete(document)
            db.commit()
            return True
        return False


document_service = DocumentService()
